#!/usr/bin/env python3
"""
Convert a markdown document into a DocX.
Output file is saved to the same directory as input.

Usage:
ONYX % python3 onyx/md_to_docx.py path/to/markdown
"""

import argparse
import os
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create the w:hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a new run element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Add style to the text (optional)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)

    # Add underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    # Add text to the run
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def markdown_to_docx(markdown_file):
    """Convert markdown file to docx."""
    # Generate output file path by replacing .md extension with .docx
    output_file = os.path.splitext(markdown_file)[0] + ".docx"

    # Read markdown file
    with open(markdown_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(
        md_content, extensions=["tables", "fenced_code", "nl2br"]
    )

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new Word document
    doc = Document()

    # Set page margins to 0.7 inches on left and right
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        # Keep default top and bottom margins
        # section.top_margin = Inches(1.0)
        # section.bottom_margin = Inches(1.0)

    # Set default document style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Configure heading styles
    for i in range(1, 7):
        style_name = f"Heading {i}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(18 - (i * 1.5))  # Decreasing font size for each level
            font.bold = True
            if i <= 2:  # Make Heading 1 and 2 stand out with color
                font.color.rgb = RGBColor(0x35, 0x53, 0x9C)  # Dark blue color

    # Setup list styles properly
    if "List Bullet" in doc.styles:
        list_style = doc.styles["List Bullet"]
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)

    if "List Number" in doc.styles:
        list_style = doc.styles["List Number"]
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)

    # Process the HTML elements recursively
    process_elements(soup, doc)

    # Save the document
    doc.save(output_file)
    print(f"Conversion complete. Output saved to {output_file}")


def process_elements(element, doc, level=0):
    """Process HTML elements and add them to the Word document."""
    if element.name == "h1":
        para = doc.add_paragraph(element.get_text(), style="Heading 1")
        para.paragraph_format.space_after = Pt(12)
    elif element.name == "h2":
        para = doc.add_paragraph(element.get_text(), style="Heading 2")
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
    elif element.name == "h3":
        para = doc.add_paragraph(element.get_text(), style="Heading 3")
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif element.name == "h4":
        para = doc.add_paragraph(element.get_text(), style="Heading 4")
    elif element.name == "h5":
        para = doc.add_paragraph(element.get_text(), style="Heading 5")
    elif element.name == "h6":
        para = doc.add_paragraph(element.get_text(), style="Heading 6")
    elif element.name == "p":
        para = doc.add_paragraph()
        process_inline_elements(element, para)
        para.paragraph_format.space_after = Pt(10)
    elif element.name == "ul":
        process_list(element, doc, is_numbered=False, level=level)
        return  # Lists handled in process_list
    elif element.name == "ol":
        process_list(element, doc, is_numbered=True, level=level)
        return  # Lists handled in process_list
    elif element.name == "table":
        process_table(element, doc)
        return  # Tables handled in process_table
    elif element.name == "pre":
        # Handle code blocks
        code = element.get_text()
        para = doc.add_paragraph()
        code_run = para.add_run(code)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(9)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        return

    # Process child elements
    if hasattr(element, "children"):
        for child in element.children:
            if child.name is not None:  # Skip NavigableString
                process_elements(child, doc, level)


def process_inline_elements(element, paragraph):
    """Process inline elements within a paragraph."""
    for child in element.children:
        if child.name is None:  # NavigableString
            run = paragraph.add_run(child.string)
        elif child.name == "strong" or child.name == "b":
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == "em" or child.name == "i":
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Courier New"
        elif child.name == "a":
            text = child.get_text()
            url = child.get("href", "")
            add_hyperlink(paragraph, text, url)
        elif child.name == "br":
            paragraph.add_run("\n")
        else:
            # Handle other inline elements
            run = paragraph.add_run(child.get_text())


def process_list(element, doc, is_numbered=False, level=0):
    """Process ordered and unordered lists."""
    # Check if this is a dashed list in the original Markdown
    is_dashed = False
    if (
        not is_numbered
        and element.find("li")
        and element.find("li").get_text().startswith("-")
    ):
        is_dashed = True

    for li in element.find_all("li", recursive=False):
        # Create paragraph with appropriate style
        para = doc.add_paragraph()
        if is_numbered:
            para.style = "List Number"
        else:
            para.style = "List Bullet"

        # Set appropriate indentation based on nesting level
        para.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
        para.paragraph_format.first_line_indent = Inches(-0.25)

        # For dashed lists (like in your business plan), use dash instead of bullet
        if is_dashed:
            # Clear any existing list formatting
            para._element.get_or_add_pPr().numPr = None
            # Add the dash manually
            para.add_run("- ")

        # Add the text
        has_processed_content = False
        for child in li.children:
            if child.name is None:  # NavigableString
                # Skip if it's just the dash character for dashed lists
                content = child.string
                if is_dashed and content.startswith("-") and not has_processed_content:
                    content = content[1:].lstrip()
                para.add_run(content)
                has_processed_content = True
            elif child.name in ["ul", "ol"]:
                # Recursive call for nested lists
                process_list(child, doc, child.name == "ol", level + 1)
            else:
                # Process other inline elements
                process_inline_elements(child, para)
                has_processed_content = True


def process_table(element, doc):
    """Process tables."""
    rows = element.find_all("tr")
    if not rows:
        return

    # Count cells in the first row to determine columns
    first_row = rows[0]
    cols = len(first_row.find_all(["th", "td"]))

    # Create the table
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    # Process each row
    for i, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for j, cell in enumerate(cells):
            if j < cols:  # Ensure we don't exceed the table dimensions
                # Add content to cell
                table_cell = table.cell(i, j)
                para = table_cell.paragraphs[0]
                process_inline_elements(cell, para)

                # Make header cells bold
                if cell.name == "th" or i == 0:  # First row is often headers
                    for run in para.runs:
                        run.bold = True

    # Add spacing after table
    doc.add_paragraph()


# Add this import
from docx.shared import RGBColor


def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX")
    parser.add_argument("input", help="Input markdown file")
    args = parser.parse_args()

    markdown_to_docx(args.input)


if __name__ == "__main__":
    main()
